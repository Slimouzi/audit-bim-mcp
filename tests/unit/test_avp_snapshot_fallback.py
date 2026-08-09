"""Pack AVP **sans sources I3F** : extraction depuis le snapshot d'audit.

Couvre l'exigence produit : si les fichiers I3F sont absents, générer SHAB,
Zones/Espaces, Enveloppe et Menuiseries depuis ``AuditResult.snapshot`` — et
interdire un livrable client réduit au seul bandeau (QA gate).
"""

from __future__ import annotations

import openpyxl
import pytest

from audit_bim.audit.engine import AuditResult
from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.profiles.i3f.tools_reporting import generate_avp_i3f_pack as tr_generate_avp_i3f_pack
from audit_bim.reporting import avp_i3f
from audit_bim.reporting.avp_i3f import AvpQaError, write_avp_i3f_report_pack
from audit_bim.reporting.avp_sources import AvpSources, MultiSheetSource, SheetGrid
from audit_bim.requirements.models import BIMPhase, RequirementsCatalog


def _catalog() -> RequirementsCatalog:
    return RequirementsCatalog(
        cch_version="3.6",
        cch_source_pdf="/tmp/cch.pdf",
        data_spec_source="/tmp/data.xlsx",
        naming_spec_source="/tmp/naming.xlsx",
        properties=[],
        naming_rules=[],
        storey_names=[],
        zone_specs=[],
        room_specs=[],
    )


def _synthetic_result() -> AuditResult:
    """Snapshot minimal : 1 mur d'enveloppe (layer cible, casse/accents/espaces
    variés), 1 espace LongName='' / Name='CHAMBRE', 1 fenêtre — surfaces
    uniquement en « Superficie calculée » (pas de BaseQuantities)."""
    wall = {
        "uuid": "W1",
        "type": "IfcWall",
        "name": "Mur ext 1",
        # Variante tolérante du layer cible « MURS - Extérieurs périphériques.Exnd ».
        "layers": [{"name": "murs  -  Exterieurs Periphériques.exnd"}],
        "property_sets": [
            {
                "name": "Pset_WallCommon",
                "properties": [{"definition": {"name": "Superficie calculée"}, "value": 42.5}],
            }
        ],
    }
    # Mur-rideau : NE DOIT PAS être compté dans l'enveloppe (décision explicite).
    curtain = {
        "uuid": "CW1",
        "type": "IfcCurtainWall",
        "name": "Façade vitrée",
        "layers": [{"name": "MURS - Extérieurs périphériques.Exnd"}],
        "property_sets": [
            {
                "name": "Pset_CurtainWallCommon",
                "properties": [{"definition": {"name": "Superficie calculée"}, "value": 999.0}],
            }
        ],
    }
    window = {
        "uuid": "WIN1",
        "type": "IfcWindow",
        "name": "F25",
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [
                    {"definition": {"name": "Width"}, "value": 0.6},
                    {"definition": {"name": "Height"}, "value": 1.3},
                ],
            }
        ],
    }
    space = {
        "uuid": "S1",
        "type": "IfcSpace",
        "name": "CHAMBRE",
        "longname": "",  # vide → le libellé exporté doit reprendre Name
        "property_sets": [
            {
                "name": "Pset_SpaceCommon",
                "properties": [{"definition": {"name": "Superficie calculée"}, "value": 12.98}],
            }
        ],
    }
    snap = ModelSnapshot(
        project={"name": "Programme"},
        model={"name": "M.ifc"},
        spaces=[space],
        elements=[wall, curtain, window],
    ).index()
    return AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])


def _all_text(path) -> str:
    wb = openpyxl.load_workbook(path, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for c in row:
                if c is not None:
                    parts.append(str(c))
    wb.close()
    return "\n".join(parts)


# ── Extraction snapshot (source absente) ────────────────────────────────


def test_snapshot_fallback_fills_annexes(tmp_path):
    result = _synthetic_result()
    # sources=None → aucune source I3F ; l'extraction snapshot doit prendre
    # le relais (sinon la QA gate lèverait).
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )

    shab = _all_text(pack.shab_xlsx)
    zones = _all_text(pack.zones_espaces_xlsx)
    env = _all_text(pack.enveloppe_xlsx)

    # SHAB / Zones : l'espace apparaît, libellé repris de Name (LongName vide).
    assert "CHAMBRE" in shab
    assert "12.98" in shab
    assert "CHAMBRE" in zones
    # Surface issue du flux maquette / IFC OpenShell, sans colonne Solibri.
    assert "Surface IFC OpenShell" in shab
    assert "Surface Solibri" not in shab

    # Enveloppe : le mur cible apparaît avec sa surface et la source tracée.
    assert "Mur ext 1" in env
    assert "42.5" in env
    assert "Archicad BQ NetSideArea" in env
    assert "Surface IFC OpenShell" in env


def _duplex_result() -> AuditResult:
    """Duplex : une zone « Logement Duplex A101 » traversant deux étages
    (R+1 / R+2), avec un séjour au R+1 et une chambre au R+2."""
    sejour = {
        "uuid": "SP-SEJOUR",
        "type": "IfcSpace",
        "name": "SEJOUR",
        "longname": "Séjour",
        "storey": {"uuid": "ST1", "name": "R+1"},
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 24.0}],
            }
        ],
    }
    chambre = {
        "uuid": "SP-CH",
        "type": "IfcSpace",
        "name": "CHAMBRE 01",
        "longname": "",  # vide → libellé repris de Name
        "storey": {"uuid": "ST2", "name": "R+2"},
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 12.0}],
            }
        ],
    }
    storeys = [{"uuid": "ST1", "name": "R+1"}, {"uuid": "ST2", "name": "R+2"}]
    zone = {
        "uuid": "Z-A101",
        "type": "IfcZone",
        "name": "Logement Duplex A101",
        "spaces": ["SP-SEJOUR", "SP-CH"],  # zone traversant les 2 étages
    }
    snap = ModelSnapshot(
        project={"name": "Programme"},
        model={"name": "M.ifc"},
        storeys=storeys,
        spaces=[sejour, chambre],
        zones=[zone],
    ).index()
    return AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])


def test_shab_export_has_zone_and_storey_columns(tmp_path):
    result = _duplex_result()
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    wb = openpyxl.load_workbook(pack.shab_xlsx, data_only=True)
    ws = wb["TDB 2022 01.3 - Export Zones..."]
    # Localise l'en-tête et vérifie la présence des colonnes Zone + Étage.
    header = None
    for row in ws.iter_rows(values_only=True):
        if row and "Composant" in [c for c in row if c is not None]:
            header = [c for c in row]
            break
    assert header is not None
    assert "Nom Zone" in header
    assert "Étage" in header
    zone_col = header.index("Nom Zone")
    storey_col = header.index("Étage")
    piece_col = header.index("Pièce")

    # Collecte les lignes de données (espaces).
    data = {}
    for row in ws.iter_rows(values_only=True):
        if row and row[piece_col] in ("Séjour", "CHAMBRE 01"):
            data[row[piece_col]] = (row[zone_col], row[storey_col])
    wb.close()

    # Les deux pièces portent la même zone (duplex) et leur étage respectif.
    assert data["Séjour"] == ("Logement Duplex A101", "R+1")
    assert data["CHAMBRE 01"] == ("Logement Duplex A101", "R+2")


def test_zones_espaces_first_tab_has_ifczone_and_storey(tmp_path):
    """Le 1er onglet de « EXPORT ZONES ET ESPACES » liste les IfcZone avec
    leur(s) étage(s) (union des étages des pièces — duplex géré)."""
    result = _duplex_result()
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    wb = openpyxl.load_workbook(pack.zones_espaces_xlsx, data_only=True)
    first = wb["TDB 2022 01.3 - Export Zones..."]  # onglet détail MOA
    rows = [tuple(r) for r in first.iter_rows(values_only=True)]
    flat = "\n".join(str(c) for r in rows for c in r if c is not None)
    wb.close()

    # En-tête Zone + Étage présent dans l'onglet détail.
    assert "Nom Zone" in flat
    assert "Étage" in flat
    # La zone duplex et les deux étages de ses pièces apparaissent.
    assert "Logement Duplex A101" in flat
    assert "R+1" in flat
    assert "R+2" in flat


def test_zones_espaces_uses_snapshot_instead_of_source_with_snapshot(tmp_path):
    """Sources I3F présentes + snapshot : l'export métier vient de la maquette."""
    from audit_bim.reporting.avp_sources import AvpSources, MultiSheetSource, SheetGrid

    result = _duplex_result()
    # Source I3F minimale avec une valeur qui ne doit pas être recopiée.
    src_grid = SheetGrid(
        title="TDB 2022 01.3 - Export Zones",
        rows=[["Composant", "Nom Zone"], ["Zone", "SOURCE-A-IGNORER"]],
    )
    sources = AvpSources(zones_espaces=MultiSheetSource(grids=[src_grid]))

    pack = write_avp_i3f_report_pack(
        result,
        tmp_path / "out",
        sources=sources,
        project_name="X",
        project_code="Y",
        export_pdf=False,
    )
    wb = openpyxl.load_workbook(pack.zones_espaces_xlsx, data_only=True)
    titles = wb.sheetnames
    flat = "\n".join(
        str(c)
        for ws in wb.worksheets
        for r in ws.iter_rows(values_only=True)
        for c in r
        if c is not None
    )
    wb.close()
    # La source n'est pas ajoutée : les onglets exportés sont ceux du snapshot.
    assert titles == ["Feuil2", "TDB 2022 01.3 - Export Zones...", "Feuil1"]
    assert "Nom Zone" in flat
    assert "Logement Duplex A101" in flat
    assert "SOURCE-A-IGNORER" not in flat


def test_generated_snapshot_reports_replace_solibri_columns(tmp_path):
    result = _duplex_result()
    sources = AvpSources(
        shab=MultiSheetSource(
            grids=[
                SheetGrid(
                    title="TDB source",
                    rows=[
                        ["Composant", "Surface Solibri"],
                        ["IfcSpace", 9999.0],
                    ],
                )
            ]
        )
    )
    pack = write_avp_i3f_report_pack(
        result,
        tmp_path / "out",
        sources=sources,
        project_name="X",
        project_code="Y",
        export_pdf=False,
    )
    txt = _all_text(pack.shab_xlsx)
    assert "Surface IFC OpenShell" in txt
    # La provenance se lit à l'emplacement de la valeur, plus dans une colonne
    # « Source quantité » ajoutée au tableau client (doctrine #210).
    assert "Source quantité" not in txt
    assert "Surface Solibri" not in txt
    assert "9999" not in txt


def test_shab_space_multiple_storeys_joined(tmp_path):
    """Un espace rattaché à deux étages (cas duplex au niveau pièce) →
    les deux étages sont listés (séparés par « / »)."""
    space = {
        "uuid": "SP-DUP",
        "type": "IfcSpace",
        "name": "SEJOUR DUPLEX",
        "storeys": [{"uuid": "ST1", "name": "R+1"}, {"uuid": "ST2", "name": "R+2"}],
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 40.0}],
            }
        ],
    }
    snap = ModelSnapshot(
        project={"name": "P"},
        model={"name": "M.ifc"},
        storeys=[{"uuid": "ST1", "name": "R+1"}, {"uuid": "ST2", "name": "R+2"}],
        spaces=[space],
    ).index()
    result = AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    txt = _all_text(pack.shab_xlsx)
    assert "R+1 / R+2" in txt


def test_envelope_excludes_curtain_wall(tmp_path):
    result = _synthetic_result()
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    env = _all_text(pack.enveloppe_xlsx)
    # Le mur-rideau (999) ne doit pas être compté dans l'enveloppe.
    assert "999" not in env
    assert "Façade vitrée" not in env


def _real_layer_wall_result() -> AuditResult:
    """Snapshot avec un mur d'enveloppe portant le nom de calque **réel** d'un
    export ArchiCAD I3F (maquette 250613_MN_BAT) : préfixe de code chantier
    « 221 - » et suffixe de vue « .Exndo »."""
    wall = {
        "uuid": "W1",
        "type": "IfcWall",
        "name": "Mur péri 221",
        "layers": [{"name": "221 - MURS - Extérieurs périphériques.Exndo"}],
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetSideArea"}, "value": 30.0}],
            }
        ],
    }
    window = {
        "uuid": "WIN1",
        "type": "IfcWindow",
        "name": "F01",
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [
                    {"definition": {"name": "Width"}, "value": 1.0},
                    {"definition": {"name": "Height"}, "value": 1.2},
                ],
            }
        ],
    }
    space = {
        "uuid": "S1",
        "type": "IfcSpace",
        "name": "SEJOUR",
        "longname": "SEJOUR",
        "property_sets": [
            {
                "name": "Pset_SpaceCommon",
                "properties": [{"definition": {"name": "Superficie calculée"}, "value": 20.0}],
            }
        ],
    }
    snap = ModelSnapshot(
        project={"name": "0546L"},
        model={"name": "250613_MN_BAT.ifc"},
        spaces=[space],
        elements=[wall, window],
    ).index()
    return AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])


def test_envelope_recognizes_real_archicad_layer_name(tmp_path):
    """Régression maquette réelle 250613_MN_BAT : un calque
    « 221 - MURS - Extérieurs périphériques.Exndo » (préfixe de code + suffixe
    « .Exndo ») doit être reconnu comme enveloppe — un match exact le ratait et
    l'annexe Enveloppe sortait vide malgré des murs présents."""
    from audit_bim.reporting.avp_snapshot import _envelope_layer_name, count_envelope_walls

    result = _real_layer_wall_result()
    snap = result.snapshot
    wall = snap.of_class("IfcWall")[0]
    # Le vrai nom de calque est reconnu (et remonté tel quel pour l'affichage).
    assert _envelope_layer_name(wall) == "221 - MURS - Extérieurs périphériques.Exndo"
    assert count_envelope_walls(snap) == 1

    pack = write_avp_i3f_report_pack(
        result,
        tmp_path / "out",
        sources=None,
        project_name="0546L",
        project_code="0546L",
        export_pdf=False,
    )
    env = _all_text(pack.enveloppe_xlsx)
    assert "Mur péri 221" in env  # annexe Enveloppe NON vide, type repris par repli Name
    assert "Archicad BQ NetSideArea" in env  # colonnes MOA du repli snapshot
    assert "30" in env  # NetSideArea remontée


def test_menuiseries_from_snapshot(tmp_path):
    result = _synthetic_result()
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    men = _all_text(pack.menuiseries_xlsx)
    assert "F25" in men  # fenêtre extraite
    # Surface = L×H (0.6×1.3 = 0.78) tracée.
    assert "0.78" in men


# ── QA gate anti-livrable vide ──────────────────────────────────────────


def test_qa_gate_raises_when_extraction_empty(tmp_path, monkeypatch):
    """Snapshot avec espaces/murs mais extraction vide (forcée) → AvpQaError."""
    result = _synthetic_result()
    # On neutralise l'extraction snapshot : les annexes sortiront vides
    # alors que le snapshot contient des données → la QA gate doit lever.
    monkeypatch.setattr(avp_i3f, "build_sources_from_snapshot", lambda snap: AvpSources())

    with pytest.raises(AvpQaError) as exc:
        write_avp_i3f_report_pack(
            result,
            tmp_path / "out",
            sources=None,
            project_name="X",
            project_code="Y",
            export_pdf=False,
        )
    # Les 3 exports concernés sont signalés.
    assert set(exc.value.empty) >= {"SHAB", "Zones/Espaces", "Enveloppe"}


def test_pack_without_snapshot_no_gate(tmp_path):
    """result=None (pas de snapshot) → la QA gate ne se déclenche pas
    (rien d'exploitable à comparer)."""
    pack = write_avp_i3f_report_pack(
        None, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    # Génère quand même le pack (bandeau + NOT_AVAILABLE), sans lever.
    assert pack.shab_xlsx.exists()


def test_snapshot_param_activates_fallback_without_result(tmp_path):
    """P1 : snapshot fourni explicitement (audit non lancé, result=None) →
    le repli maquette s'active, les annexes sont remplies (pas de bandeau
    seul, pas d'AvpQaError)."""
    result = _synthetic_result()
    snap = result.snapshot
    pack = write_avp_i3f_report_pack(
        None,  # pas d'AuditResult (verify_active_model seul)
        tmp_path / "out",
        sources=None,
        snapshot=snap,
        project_name="X",
        project_code="Y",
        export_pdf=False,
    )
    assert "CHAMBRE" in _all_text(pack.shab_xlsx)
    assert "Mur ext 1" in _all_text(pack.enveloppe_xlsx)


def test_tool_passes_state_snapshot_without_audit(tmp_path, monkeypatch):
    """P1 (bout en bout) : après verify_active_model seul (_State.snapshot
    posé, _State.result None), generate_avp_i3f_pack remplit les annexes."""
    from audit_bim.mcp.session import _Session, current_session

    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    # Test HERMÉTIQUE : dossier d'entrée vide, et auto-calcul enveloppe coupé.
    # Sans cela, il dépendait des fichiers présents dans le workspace (un
    # ``*_envelope.json`` traînant le faisait passer en local et échouer en CI).
    # L'auto-résolution a ses propres tests : test_avp_pack_autocompute.py.
    entree = tmp_path / "in_vide"
    entree.mkdir()
    monkeypatch.setenv("AUDIT_INPUT_DIR", str(entree))
    sess = _Session()
    sess.snapshot = _synthetic_result().snapshot  # snapshot chargé, pas d'audit
    token = current_session.set(sess)
    try:
        res = tr_generate_avp_i3f_pack(
            project_name="X",
            project_code="Y",
            phase="AVP",
            auditor="AMO BIM",
            export_pdf=False,
            auto_compute_envelope=False,
        )
    finally:
        current_session.reset(token)
    assert res.get("status") != "error"
    shab = next(p for p in res["paths"] if "SHAB" in p)
    assert "CHAMBRE" in _all_text(shab)


def test_storey_from_structure_tree(tmp_path):
    """P2 : l'étage n'est présent que dans structure_tree → l'export SHAB
    doit tout de même remplir la colonne Étage."""
    space = {
        "uuid": "SP1",
        "type": "IfcSpace",
        "name": "CHAMBRE",
        "longname": "",
        # PAS de storey en attribut plat.
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 12.0}],
            }
        ],
    }
    tree = [
        {
            "type": "IfcBuildingStorey",
            "uuid": "ST1",
            "name": "R+3",
            "children": [
                {
                    "type": "IfcZone",
                    "uuid": "Z1",
                    "name": "Logement Z",
                    "children": [{"type": "IfcSpace", "uuid": "SP1", "name": "CHAMBRE"}],
                }
            ],
        }
    ]
    snap = ModelSnapshot(
        project={"name": "P"}, model={"name": "M.ifc"}, spaces=[space], structure_tree=tree
    ).index()
    result = AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    txt = _all_text(pack.shab_xlsx)
    assert "R+3" in txt  # étage résolu via structure_tree
    assert "Logement Z" in txt  # zone résolue via structure_tree


def test_zones_grid_members_from_structure_tree(tmp_path):
    """P1 : /zone sans liste ``spaces`` mais structure_tree contient
    Zone → Space → l'onglet Zones remplit Étage(s) / Nombre de pièces /
    Surface."""
    sejour = {
        "uuid": "SP-A",
        "type": "IfcSpace",
        "name": "SEJOUR",
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 20.0}],
            }
        ],
    }
    chambre = {
        "uuid": "SP-B",
        "type": "IfcSpace",
        "name": "CHAMBRE",
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 10.0}],
            }
        ],
    }
    # La zone NE porte PAS de liste spaces ; l'arborescence, si.
    zone = {"uuid": "Z1", "type": "IfcZone", "name": "Logement Z"}
    tree = [
        {
            "type": "IfcBuildingStorey",
            "uuid": "ST1",
            "name": "R+1",
            "children": [
                {
                    "type": "IfcZone",
                    "uuid": "Z1",
                    "name": "Logement Z",
                    "children": [
                        {"type": "IfcSpace", "uuid": "SP-A", "name": "SEJOUR"},
                        {"type": "IfcSpace", "uuid": "SP-B", "name": "CHAMBRE"},
                    ],
                }
            ],
        }
    ]
    snap = ModelSnapshot(
        project={"name": "P"},
        model={"name": "M.ifc"},
        spaces=[sejour, chambre],
        zones=[zone],
        structure_tree=tree,
    ).index()
    result = AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    wb = openpyxl.load_workbook(pack.zones_espaces_xlsx, data_only=True)
    first = wb["TDB 2022 01.3 - Export Zones..."]  # onglet détail MOA
    rows = [tuple(r) for r in first.iter_rows(values_only=True)]
    flat = "\n".join(str(c) for r in rows for c in r if c is not None)
    wb.close()
    assert "Logement Z" in flat
    assert "R+1" in flat  # étage retrouvé via structure_tree
    assert "20" in flat and "10" in flat  # surfaces des pièces rattachées
    assert "CHAMBRE" in flat and "SEJOUR" in flat


def test_menuiseries_standardcase_ifc4(tmp_path):
    """P2 : un IFC4 avec uniquement des …StandardCase remplit Menuiseries
    (builder) et n'échappe pas à la QA gate."""
    win = {
        "uuid": "W-SC",
        "type": "IfcWindowStandardCase",
        "name": "F-SC",
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [
                    {"definition": {"name": "Width"}, "value": 1.0},
                    {"definition": {"name": "Height"}, "value": 1.2},
                ],
            }
        ],
    }
    door = {"uuid": "D-SC", "type": "IfcDoorStandardCase", "name": "P-SC"}
    space = {
        "uuid": "S1",
        "type": "IfcSpace",
        "name": "CHAMBRE",
        "property_sets": [
            {
                "name": "BaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": 12.0}],
            }
        ],
    }
    snap = ModelSnapshot(
        project={"name": "P"}, model={"name": "M.ifc"}, spaces=[space], elements=[win, door]
    ).index()
    result = AuditResult(phase=BIMPhase.AVP, catalog=_catalog(), snapshot=snap, findings=[])
    # Ne doit PAS lever (Menuiseries remplies via StandardCase).
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    men = _all_text(pack.menuiseries_xlsx)
    assert "F-SC" in men
    # La PORTE ne doit PLUS apparaître : ce livrable a pour gabarit
    # « TDB 2022 05.1 - Fenêtres Ok », qui ne contient que des fenêtres. Les
    # portes relèvent d'un livrable distinct, non encore spécifié côté client.
    assert "P-SC" not in men


def test_docx_ecarts_uses_superficie_calculee_fallback(tmp_path):
    """P2 : la section Écarts du DOCX calcule la SHAB snapshot avec le repli
    « Superficie calculée » (espaces sans BaseQuantities)."""
    from docx import Document

    result = _synthetic_result()  # espace surface uniquement en Superficie calculée
    pack = write_avp_i3f_report_pack(
        result, tmp_path / "out", sources=None, project_name="X", project_code="Y", export_pdf=False
    )
    doc = Document(str(pack.analyse_docx))
    txt = "\n".join(p.text for p in doc.paragraphs)
    txt += "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    # SHAB snapshot = 12.98 (Superficie calculée) doit apparaître, pas NOT_AVAILABLE.
    assert "12.98" in txt


def test_qa_gate_blocks_empty_menuiseries(tmp_path, monkeypatch):
    """P3 : IfcWindow présent mais export Menuiseries vide → AvpQaError."""
    result = _synthetic_result()  # contient une IfcWindow
    monkeypatch.setattr(avp_i3f, "build_sources_from_snapshot", lambda snap: AvpSources())
    with pytest.raises(AvpQaError) as exc:
        write_avp_i3f_report_pack(
            result,
            tmp_path / "out",
            sources=None,
            project_name="X",
            project_code="Y",
            export_pdf=False,
        )
    assert "Menuiseries" in exc.value.empty


def test_tool_returns_error_status_on_empty(tmp_path, monkeypatch):
    """Le tool MCP renvoie un statut d'erreur explicite (pas un fichier vide)
    quand la QA gate échoue."""
    from audit_bim.mcp.session import _Session, current_session

    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    monkeypatch.setattr(avp_i3f, "build_sources_from_snapshot", lambda snap: AvpSources())

    sess = _Session()
    sess.result = _synthetic_result()
    token = current_session.set(sess)
    try:
        res = tr_generate_avp_i3f_pack(
            project_name="X", project_code="Y", phase="AVP", auditor="AMO BIM", export_pdf=False
        )
    finally:
        current_session.reset(token)

    assert res.get("status") == "error"
    assert res.get("error") == "empty_deliverable"
    assert "SHAB" in res["empty_deliverables"]
