"""Lot 4 — restitution des quantités calculées : colonne « Source quantité »,
note méthodo, statut partial_computed, couverture dans le .docx."""

from __future__ import annotations

import openpyxl
from docx import Document

from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.reporting.avp_availability import inspect_avp_report_availability
from audit_bim.reporting.avp_i3f import write_avp_i3f_report_pack
from audit_bim.reporting.avp_snapshot import build_shab_from_snapshot


def _space(uuid, name, value, *, computed):
    el = {
        "uuid": uuid,
        "type": "IfcSpace",
        "name": name,
        "property_sets": [
            {
                "name": "Qto_SpaceBaseQuantities",
                "properties": [{"definition": {"name": "NetFloorArea"}, "value": value}],
            }
        ],
    }
    if computed:
        el["computed_base_quantities"] = [
            {
                "quantity": "NetFloorArea",
                "qto": "Qto_SpaceBaseQuantities",
                "value": value,
                "unit": "m2",
                "method": "ifcopenshell_geometry",
                "status": "computed",
                "source": "computed_ifcopenshell",
            }
        ]
    return el


def _snapshot():
    # S1 = surface calculée (fusion) ; S2 = surface native BIMData ; 1 zone (relation).
    return ModelSnapshot(
        project={"name": "DIEPPE"},
        model={"name": "M.ifc"},
        spaces=[
            _space("S1", "CH1", 12.3, computed=True),
            _space("S2", "CH2", 99.0, computed=False),
        ],
        zones=[{"uuid": "Z1", "type": "IfcZone", "name": "0546L-1", "spaces": ["S1", "S2"]}],
    ).index()


def _all_text(path):
    wb = openpyxl.load_workbook(path, data_only=True)
    parts = [
        str(c)
        for ws in wb.worksheets
        for row in ws.iter_rows(values_only=True)
        for c in row
        if c is not None
    ]
    wb.close()
    return "\n".join(parts)


def _docx_text(path):
    d = Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# ── colonne « Source quantité » dans les builders ───────────────────────


def test_la_colonne_de_mesure_distingue_le_calcule_du_natif():
    """La provenance se lit à **l'emplacement de la valeur**, non dans une
    colonne « Source quantité » ajoutée au tableau client (doctrine #210)."""
    ms, _total = build_shab_from_snapshot(_snapshot())
    rows = next(g.rows for g in ms.grids if g.title.startswith("TDB 2022 01.3"))
    header = rows[0]
    assert "Source quantité" not in header
    calc_idx = header.index("Surface IFC OpenShell")
    natif_idx = header.index("Surface Nette (Qté de Base)")
    piece_idx = header.index("Pièce")
    par_piece = {r[piece_idx]: (r[calc_idx], r[natif_idx]) for r in rows[1:] if any(r)}
    assert par_piece["CH1"][0] is not None and par_piece["CH1"][1] is None
    assert par_piece["CH2"][1] is not None and par_piece["CH2"][0] is None


# ── availability : statut partial_computed ──────────────────────────────


def test_availability_partial_computed_when_computed_assisted():
    avails = {a.key: a for a in inspect_avp_report_availability(_snapshot(), has_audit_result=True)}
    shab = avails["shab_maquette"]
    assert shab.can_generate is True
    assert shab.status == "partial_computed"
    assert shab.computed_assisted is True
    # Le caractère calculé / non contractuel est affiché clairement.
    assert "calculées" in shab.next_action.lower()


def test_gaps_not_hidden_when_quantity_absent():
    # Menuiseries présentes mais SANS Width/Height (ni natif ni calculé) → gap
    # affiché dans missing_data, jamais masqué.
    win = {"uuid": "W1", "type": "IfcWindow", "name": "F1", "property_sets": []}
    snap = ModelSnapshot(elements=[win]).index()
    men = {a.key: a for a in inspect_avp_report_availability(snap)}["menuiseries"]
    assert any("Width" in m or "Height" in m for m in men.missing_data)


def test_availability_stays_partial_without_computed():
    # Snapshot avec NetFloorArea native (aucune computed) → partial classique.
    native_only = ModelSnapshot(
        spaces=[_space("S2", "CH2", 99.0, computed=False)],
        zones=[{"uuid": "Z1", "type": "IfcZone", "name": "0546L-1", "spaces": ["S2"]}],
    ).index()
    shab = {a.key: a for a in inspect_avp_report_availability(native_only)}["shab_maquette"]
    assert shab.status == "partial" and shab.computed_assisted is False


# ── pack : note méthodo (xlsx) + couverture (docx) ──────────────────────


def test_pack_shab_carries_methodo_note_and_computed_cell(tmp_path):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", snapshot=_snapshot(), export_pdf=False)
    text = _all_text(pack.shab_xlsx)
    # Le caractère calculé se lit désormais dans l'en-tête de la colonne qui
    # porte la valeur, et non dans un libellé répété en bout de ligne.
    assert "Surface IFC OpenShell" in text
    # La note méthodo dit le RÔLE de chaque colonne, et à quelle condition une
    # quantité n'est pas contractuelle : lorsque seule la colonne calculée est
    # renseignée. Elle n'affirme plus qu'une quantité est « native OU calculée »
    # — les deux coexistent dès qu'une valeur de comparaison existe.
    assert "Provenance des quantités" in text
    assert "NON contractuelle" in text
    assert "peuvent coexister" in text
    # Orthographe canonique des livrables : « IFC OpenShell », pas le nom de la
    # bibliothèque. Le texte client ne cite plus « IfcOpenShell ».
    assert "IFC OpenShell" in text


def test_docx_reports_computed_coverage(tmp_path):
    snap = _snapshot()
    snap.computed_coverage = {
        "n_merged": 7,
        "n_gap_kept": 3,
        "n_skipped_status": 2,
        "n_unknown_uuid": 1,
    }
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", snapshot=snap, export_pdf=False)
    txt = _docx_text(pack.analyse_docx)
    assert "Quantités calculées (IfcOpenShell) — couverture" in txt
    # Les 4 compteurs demandés apparaissent.
    for n in ("7", "3", "2", "1"):
        assert n in txt
